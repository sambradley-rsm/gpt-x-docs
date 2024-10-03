import os
import re
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Specify the folder for Markdown and Word documents
MD_FOLDER = '../docs/markdown'
DOCX_FOLDER = '../docs/docx'

def convert_md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()

    # Split the text by lines
    lines = markdown_text.splitlines()

    # Create a new Word document
    doc = Document()

    # Keep track of whether we are in a list
    in_list = False

    # Parse the markdown text and apply styles for headings and paragraphs
    for line in lines:
        line = line.strip()

        if line.startswith('# '):
            paragraph = doc.add_heading(level=1)  # Heading 1
            parse_and_add_bold_and_links(paragraph, line[2:], doc)
            in_list = False  # End any list
        elif line.startswith('## '):
            paragraph = doc.add_heading(level=2)  # Heading 2
            parse_and_add_bold_and_links(paragraph, line[3:], doc)
            in_list = False  # End any list
        elif line.startswith('### '):
            paragraph = doc.add_heading(level=3)  # Heading 3
            parse_and_add_bold_and_links(paragraph, line[4:], doc)
            in_list = False  # End any list
        elif line.startswith('#### '):
            paragraph = doc.add_heading(level=4)  # Heading 4
            parse_and_add_bold_and_links(paragraph, line[5:], doc)
            in_list = False  # End any list
        elif line.startswith(('- ', '* ', '• ')):
            # Handle bullet points
            paragraph = doc.add_paragraph(style='ListBullet')
            parse_and_add_bold_and_links(paragraph, line[2:], doc)
        elif line.startswith('!['):  # Handle images
            handle_image(line, md_file, doc)
        elif re.search(r'\[.*?\]\(.*?\)\{.*md-button.*\}', line):  # Handle button-like hyperlinks
            handle_button_hyperlink(line, doc)
        elif line:
            paragraph = doc.add_paragraph()
            parse_and_add_bold_and_links(paragraph, line, doc)
            
    # Save the Word document
    doc.save(docx_file)

def handle_image(line, md_file, doc):
    # Modify the regex pattern to make the alignment part optional
    image_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)(?:\{.*align=(right|left|center).*?\})?')
    match = image_pattern.search(line)
    
    if match:
        alt_text, image_path, alignment = match.groups()
        
        # Adjust the image path relative to the project root (assume it's '../assets/')
        project_root = os.path.abspath(os.path.join(os.path.dirname(md_file), '..'))
        full_image_path = os.path.normpath(os.path.join(project_root, image_path.strip('./')))

        # Print the resolved image path to verify it
        print(f"Resolved image path: {full_image_path}")

        # Check if the file exists before trying to insert it
        if not os.path.exists(full_image_path):
            print(f"Image file not found: {full_image_path}")
            return

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        try:
            run.add_picture(full_image_path, width=Inches(2))  # Adjust the image size as needed
        except Exception as e:
            print(f"Error inserting image: {e}")
            return

        # Set alignment - default to left if not specified
        if alignment == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Default alignment


# Function to handle button-like hyperlinks
def handle_button_hyperlink(line, doc):
    button_pattern = re.compile(r'\[(.*?)\]\((.*?)\)\{.*md-button.*\}')
    match = button_pattern.search(line)
    if match:
        display_text, url = match.groups()
        paragraph = doc.add_paragraph()
        add_button_hyperlink(paragraph, display_text, url)

def add_button_hyperlink(paragraph, text, url):
    """
    Add a button-like hyperlink with white text on a dark blue background.
    """
    # Create a hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the <w:hyperlink> element and add the relationship ID
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the hyperlink
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add bold styling for the button text
    bold = OxmlElement('w:b')
    rPr.append(bold)

    # Add white text color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF')  # White text for button
    rPr.append(color)

    # Add blue background to simulate a button
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), '000080')  # Dark blue background color
    rPr.append(shading)

    # Create text for the hyperlink
    run_text = OxmlElement('w:t')
    run_text.text = text

    # Append the run properties and text to the run
    run.append(rPr)
    run.append(run_text)

    # Append the run to the hyperlink element (this makes it clickable)
    hyperlink.append(run)

    # Append the hyperlink element to the paragraph
    paragraph._element.append(hyperlink)

    # left align the paragraph to make the button look good
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Function to parse and add bold text and hyperlinks
def parse_and_add_bold_and_links(paragraph, text, doc):
    # Regex to find bold text (**text**)
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    # Regex to find hyperlinks ([text](url))
    link_pattern = re.compile(r'\[(.*?)\]\((.*?)\)')

    current_pos = 0

    while current_pos < len(text):
        # Check for bold or link matches
        bold_match = bold_pattern.search(text, current_pos)
        link_match = link_pattern.search(text, current_pos)

        next_match = None
        if bold_match and (not link_match or bold_match.start() < link_match.start()):
            next_match = bold_match
        elif link_match:
            next_match = link_match

        if not next_match:
            # If no matches, add remaining text
            paragraph.add_run(text[current_pos:])
            break

        # Add text before match
        if next_match.start() > current_pos:
            paragraph.add_run(text[current_pos:next_match.start()])

        if next_match == bold_match:
            # Handle bold text
            bold_run = paragraph.add_run(next_match.group(1))
            bold_run.bold = True
            current_pos = bold_match.end()
        elif next_match == link_match:
            # Handle links
            add_hyperlink(paragraph, next_match.group(2), next_match.group(1))
            current_pos = link_match.end()

def add_hyperlink(paragraph, url, display_text):
    """
    Add a hyperlink to a Word paragraph and style it with blue underlined text.
    :param paragraph: The paragraph object to add the hyperlink to.
    :param url: The URL for the hyperlink.
    :param display_text: The text to display for the hyperlink.
    """
    # Create a hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set the hyperlink style (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # Underline the text
    rPr.append(underline)

    run.append(rPr)

    # Add the display text for the hyperlink
    run_text = OxmlElement('w:t')
    run_text.text = display_text
    run.append(run_text)

    hyperlink.append(run)

    # Append the hyperlink element to the paragraph
    paragraph._element.append(hyperlink)

# List .md files and prompt for selection
def select_md_file():
    files = [f for f in os.listdir(MD_FOLDER) if f.endswith(".md")]
    if not files:
        print("No Markdown files found.")
        return None
    print("Select a Markdown file to convert to Word:")
    for i, file in enumerate(files, 1):
        print(f"{i}. {file}")
    
    choice = input(f"Enter the number (1-{len(files)}): ")
    try:
        selected_file = files[int(choice) - 1]
        return os.path.join(MD_FOLDER, selected_file)
    except (IndexError, ValueError):
        print("Invalid choice.")
        return None

if __name__ == "__main__":
    md_file = select_md_file()
    if md_file:
        docx_filename = os.path.basename(md_file).replace('.md', '.docx')
        docx_file = os.path.join(DOCX_FOLDER, docx_filename)
        
        # Normalize the path to ensure correct separators
        docx_file = docx_file.replace('\\', '/')  # Replace backslashes with forward slashes
        
        convert_md_to_docx(md_file, docx_file)
        print(f"Converted {md_file} to {docx_file}")
